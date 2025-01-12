import streamlit as st
import requests
import json
from io import BytesIO
import logging
import os
import fitz
import ast
import re
import uuid
import redis
from openai import AzureOpenAI
import asyncio
from concurrent.futures import ThreadPoolExecutor
import tiktoken
from pydantic import BaseModel, Field, ValidationError
from typing import List, Dict, Union
from typing import List, Dict, Union
from sentence_transformers import SentenceTransformer, util
from fuzzywuzzy import fuzz
from docx import Document
from metaphone import doublemetaphone
import re
from nltk.stem import WordNetLemmatizer

semantic_model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

llm_api_key = os.getenv("AZURE_API_KEY")
azure_llm_endpoint = os.getenv("AZURE_ENDPOINT")
llm_model = os.getenv("LLM_MODEL")
llm_api_version = "2024-10-01-preview"

redis_host = os.getenv("REDIS_HOST")
redis_port = 6379
redis_password = os.getenv("REDIS_KEY")

# Connect to Azure Redis
redis_client = redis.StrictRedis(
    host=redis_host, port=redis_port, password=redis_password, decode_responses=True
)

llm_headers = {"Content-Type": "application/json", "api-key": llm_api_key}

if "documents" not in st.session_state:
    st.session_state.documents = {}
if "removed_documents" not in st.session_state:
    st.session_state.removed_documents = []
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())


st.title("AttorneyAI")

uploaded_files = st.file_uploader(
    "Upload files less than 400 pages",
    type=["pdf", "docx", "xlsx", "pptx"],
    accept_multiple_files=True,
    help="If your question is not answered properly or there's an error, consider uploading smaller documents or splitting larger ones.",
    label_visibility="collapsed",
)


def print_extracted_details_from_redis(redis_key):
    try:
        stored_data = redis_client.get(redis_key)
        if stored_data:
            details = json.loads(stored_data)
            print(f"Details retrieved from Redis for {redis_key}: {details}")
            st.write(f"Details retrieved from Redis for {redis_key}: {details}")
        else:
            print(f"No data found in Redis for key {redis_key}")
    except Exception as e:
        print(f"Error retrieving data from Redis: {e}")


def compare_trademarks2(
    existing_trademark: List[Dict[str, Union[str, List[int]]]],
    proposed_name: str,
    proposed_class: str,
    proposed_goods_services: str,
) -> List[Dict[str, Union[str, int]]]:
    proposed_classes = [int(c.strip()) for c in proposed_class.split(",")]

    # Prepare the messages for the Azure OpenAI API
    messages = [
        {
            "role": "system",
            "content": """  
            You are a trademark attorney tasked with determining a conflict grade based on the given conditions.  
           
            Additional Instructions: 
           
            - Consider if the proposed trademark name appears anywhere within the existing trademark name, or if significant parts of the existing trademark name appear in the proposed name.  
            - Evaluate shared words between trademarks, regardless of their position.  
            - Assess phonetic similarities, including partial matches or subtle matches.  
            - Consider the overall impression created by the trademarks, including similarities in appearance, sound, pronounciation, and meaning.  
           
            Follow the conflict grading criteria as previously outlined, assigning "Name-Match" or "No-conflict" based on your analysis.  
            """,
        },
        {
            "role": "user",
            "content": f"""  
            Evaluate the potential conflict between the following existing trademarks and the proposed trademark.  
           
            Proposed Trademark:
            - Name: "{proposed_name}"  
           
            Existing Trademarks:
            - Name: "{existing_trademark['-_trademark_name']}"  
            - Status: "{existing_trademark['-_status']}"
           
            Instructions:
            1. Review the proposed and existing trademark data.  
            2. Determine if the trademarks are likely to cause confusion based on the Trademark name such as Phonetic match, Semantic similarity and String similarity.  
            3. Return the output with Conflict Grade only as 'Name-Match' or 'No-conflict', based on the reasoning.
            4. Provide reasoning for each Conflict Grade.
            5. Special Case: If the existing trademark status is "Cancelled" or "Abandoned," it will automatically be considered as Conflict Grade: No-conflict.  
           
            Output Format:
                Existing Name: Name of the existing trademark.
                Reasoning: Reasoning.
                Conflict Grade: Name-Match/No-conflict
        """,
        },
    ]

    # Initialize the Azure OpenAI client
    azure_endpoint = azure_llm_endpoint
    api_key = llm_api_key

    if not azure_endpoint or not api_key:
        raise ValueError(
            "Azure endpoint or API key is not set in environment variables."
        )

    client = AzureOpenAI(
        azure_endpoint=azure_llm_endpoint,
        api_key=llm_api_key,
        api_version="2024-10-01-preview",
    )

    # Call Azure OpenAI to get the response
    try:
        response_reasoning = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            temperature=0,
            max_tokens=500,
            top_p=1,
        )

        # Extract the content from the response
        reasoning_content = response_reasoning.choices[0].message.content
        conflict_grade = reasoning_content.split("Conflict Grade:", 1)[1].strip()
        st.write(reasoning_content)

        return conflict_grade

    except Exception as e:
        print(f"Error while calling Azure OpenAI API: {e}")
        return []


def compare_trademarks(
    existing_trademark: Dict[str, Union[str, List[int]]],
    proposed_name: str,
    proposed_class: str,
    proposed_goods_services: str,
) -> Dict[str, Union[str, int]]:
    # Convert proposed classes to a list of integers
    international_class_numbers = ast.literal_eval(
        existing_trademark["-_international_class_number"]
    )
    proposed_classes = [int(c.strip()) for c in proposed_class.split(",")]
    if not (any(cls in international_class_numbers for cls in proposed_classes)):
        conflict_grade = compare_trademarks2(existing_trademark, proposed_name, proposed_class, proposed_goods_services)
        return {
        "Trademark name": existing_trademark["-_trademark_name"],
        "Trademark -_status": existing_trademark["-_status"],
        "Trademark -_owner": existing_trademark["-_owner"],
        "Trademark class Number": existing_trademark["-_international_class_number"],
        "Trademark serial number": existing_trademark["-_serial_number"],
        "Trademark registration number": existing_trademark["-_registration_number"],
        "Trademark design phrase": existing_trademark["-_design_phrase"],
        "conflict_grade": conflict_grade,
        "reasoning": "reasoning",
    }
        
    # Helper function for semantic equivalence
    def is_semantically_equivalent(name1, name2, threshold=0.80):
        embeddings1 = semantic_model.encode(name1, convert_to_tensor=True)
        embeddings2 = semantic_model.encode(name2, convert_to_tensor=True)
        similarity_score = util.cos_sim(embeddings1, embeddings2).item()
        return similarity_score >= threshold

    # Helper function for phonetic equivalence
    def is_phonetically_equivalent(name1, name2, threshold=80):
        return fuzz.ratio(name1.lower(), name2.lower()) >= threshold

    # Helper function for phonetically equivalent words
    def first_words_phonetically_equivalent(existing_name, proposed_name, threshold=80):
        existing_words = existing_name.lower().split()
        proposed_words = proposed_name.lower().split()
        if len(existing_words) < 2 or len(proposed_words) < 2:
            return False
        return (
            fuzz.ratio(" ".join(existing_words[:2]), " ".join(proposed_words[:2]))
            >= threshold
        )

    def is_exact_match(name1: str, name2: str) -> bool:
        # Initial exact match check
        if name1.strip().lower() == name2.strip().lower():
            return True
        else:
            # Check for near-exact matches using normalized forms
            normalized_name1 = normalize_text(name1)
            normalized_name2 = normalize_text(name2)
            if normalized_name1 == normalized_name2:
                return True
            elif fuzz.ratio(normalized_name1, normalized_name2) >= 95:
                # Near-exact match, supplement with LLM
                return is_exact_match_llm(name1, name2)
            else:
                return False

    def normalize_text(text: str) -> str:
        import unicodedata
        import re

        # Normalize unicode characters
        text = unicodedata.normalize("NFKD", text)
        # Remove diacritics
        text = "".join(c for c in text if not unicodedata.combining(c))
        # Remove special characters and punctuation
        text = re.sub(r"[^\w\s]", "", text)
        # Convert to lowercase and strip whitespace
        return text.lower().strip()

    def is_exact_match_llm(name1: str, name2: str) -> bool:
        from openai import AzureOpenAI
        import os

        azure_endpoint = azure_llm_endpoint
        api_key = llm_api_key
        client = AzureOpenAI(
            azure_endpoint=azure_endpoint,
            api_key=api_key,
            api_version="2024-10-01-preview",
        )

        prompt = f"""  
            Are the following two trademark names considered exact matches, accounting for minor variations such as special characters, punctuation, or formatting? Respond with 'Yes' or 'No'.  
            
            Trademark Name 1: "{name1}"  
            Trademark Name 2: "{name2}"  
            """

        messages = [
            {
                "role": "system",
                "content": "You are a trademark expert specializing in name comparisons.",
            },
            {"role": "user", "content": prompt},
        ]

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            temperature=0.0,
            max_tokens=5,
        )

        answer = response.choices[0].message.content.strip().lower()
        return "yes" in answer.lower()

    def is_semantically_equivalents(
        name1: str, name2: str, threshold: float = 0.80
    ) -> bool:
        embeddings1 = semantic_model.encode(name1, convert_to_tensor=True)
        embeddings2 = semantic_model.encode(name2, convert_to_tensor=True)
        similarity_score = util.cos_sim(embeddings1, embeddings2).item()
        if similarity_score >= threshold:
            return True
        elif similarity_score >= (threshold - 0.1):
            # Near-threshold case, supplement with LLM
            return is_semantically_equivalent_llm(name1, name2)
        else:
            return False

    def is_semantically_equivalent_llm(name1: str, name2: str) -> bool:
        prompt = f"""  
        Are the following two trademark names semantically equivalent? Respond with 'Yes' or 'No'.  
        
        Trademark Name 1: "{name1}"  
        Trademark Name 2: "{name2}"  
        """

        azure_endpoint = azure_llm_endpoint
        api_key = llm_api_key
        client = AzureOpenAI(
            azure_endpoint=azure_endpoint,
            api_key=api_key,
            api_version="2024-10-01-preview",
        )

        messages = [
            {
                "role": "system",
                "content": "You are an expert in trademark law and semantics.",
            },
            {"role": "user", "content": prompt},
        ]

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            temperature=0.0,
            max_tokens=5,
        )

        answer = response.choices[0].message.content.strip().lower()
        return "yes" in answer.lower()

    def is_phonetically_equivalents(
        name1: str, name2: str, threshold: int = 80
    ) -> bool:

        dm_name1 = doublemetaphone(name1)
        dm_name2 = doublemetaphone(name2)
        phonetic_similarity = fuzz.ratio(dm_name1[0], dm_name2[0])
        if phonetic_similarity >= threshold:
            return True
        elif phonetic_similarity >= (threshold - 10):
            # Near-threshold case, supplement with LLM
            return is_phonetically_equivalent_llm(name1, name2)
        else:
            return False

    def is_phonetically_equivalent_llm(name1: str, name2: str) -> bool:

        prompt = f"""  
        Do the following two trademark names sound the same or very similar when spoken aloud? Consider differences in spelling but similarities in pronunciation. Respond with 'Yes' or 'No'.  
        
        Trademark Name 1: "{name1}"  
        Trademark Name 2: "{name2}"  
        """

        messages = [
            {
                "role": "system",
                "content": "You are an expert in phonetics and trademark law.",
            },
            {"role": "user", "content": prompt},
        ]

        azure_endpoint = azure_llm_endpoint
        api_key = llm_api_key
        client = AzureOpenAI(
            azure_endpoint=azure_endpoint,
            api_key=api_key,
            api_version="2024-10-01-preview",
        )

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            temperature=0.0,
            max_tokens=5,
        )

        answer = response.choices[0].message.content.strip().lower()
        return "yes" in answer.lower()

    # Condition 1A: Exact character-for-character match
    condition_1A_satisfied = (
        existing_trademark["-_trademark_name"].strip().lower()
        == proposed_name.strip().lower()
    )

    condition_1A_satisfieds = is_exact_match(
        existing_trademark["-_trademark_name"].strip().lower(),
        proposed_name.strip().lower(),
    )
    st.write(f"Exact Match: {condition_1A_satisfieds}")

    # Condition 1B: Semantically equivalent
    condition_1B_satisfied = is_semantically_equivalent(
        existing_trademark["-_trademark_name"], proposed_name
    )

    condition_1B_satisfieds = is_semantically_equivalents(
        existing_trademark["-_trademark_name"].strip().lower(),
        proposed_name.strip().lower(),
    )
    st.write(f"Semantically equivalents : {condition_1B_satisfieds}")

    # Condition 1C: Phonetically equivalent
    condition_1C_satisfied = is_phonetically_equivalent(
        existing_trademark["-_trademark_name"], proposed_name
    )

    condition_1C_satisfieds = is_phonetically_equivalents(
        existing_trademark["-_trademark_name"], proposed_name
    )
    st.write(f"Phonetically equivalents : {condition_1C_satisfieds}")

    # Condition 1D: First two or more words are phonetically equivalent
    condition_1D_satisfied = first_words_phonetically_equivalent(
        existing_trademark["-_trademark_name"], proposed_name
    )

    # Condition 1E: Proposed name is the first word of the existing trademark
    condition_1E_satisfied = (
        existing_trademark["-_trademark_name"].lower().startswith(proposed_name.lower())
    )

    # Check if any Condition 1 is satisfied
    condition_1_satisfied = any(
        [
            condition_1A_satisfied,
            condition_1B_satisfied,
            condition_1C_satisfied,
            condition_1D_satisfied,
            condition_1E_satisfied,
        ]
    )

    def target_market_and_goods_overlaps(existing_gs, proposed_gs, threshold=0.65):
        embeddings1 = semantic_model.encode(existing_gs, convert_to_tensor=True)
        embeddings2 = semantic_model.encode(proposed_gs, convert_to_tensor=True)
        similarity_score = util.cos_sim(embeddings1, embeddings2).item()
        if similarity_score >= threshold:
            return True
        elif similarity_score >= (threshold - 0.1):
            # Supplement with LLM
            return target_market_and_goods_overlap_llm(existing_gs, proposed_gs)
        else:
            # Further check using keyword overlap
            # ... Additional code
            return False

    def target_market_and_goods_overlap_llm(existing_gs: str, proposed_gs: str) -> bool:
        prompt = f"""  
            Do the goods and services described in the existing trademark and the proposed trademark overlap or target the same market? Consider the descriptions carefully. Respond with 'Yes' or 'No'.  
            
            Existing Trademark Goods/Services:  
            "{existing_gs}"  
            
            Proposed Trademark Goods/Services:  
            "{proposed_gs}"  
            """

        messages = [
            {
                "role": "system",
                "content": "You are an expert in trademark law and market analysis.",
            },
            {"role": "user", "content": prompt},
        ]

        azure_endpoint = azure_llm_endpoint
        api_key = llm_api_key
        client = AzureOpenAI(
            azure_endpoint=azure_endpoint,
            api_key=api_key,
            api_version="2024-10-01-preview",
        )

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            temperature=0.0,
            max_tokens=5,
        )

        answer = response.choices[0].message.content.strip().lower()
        return "yes" in answer.lower()

    # Condition 2: Overlap in International Class Numbers
    international_class_numbers = ast.literal_eval(
        existing_trademark["-_international_class_number"]
    )

    # Check if any class in proposed_classes is in the international_class_numbers
    condition_2_satisfied = any(
        cls in international_class_numbers for cls in proposed_classes
    )


    def normalize_text(text):

        # Replace special hyphen-like characters with a standard hyphen
        text = re.sub(r"[−–—]", "-", text)
        # Remove punctuation except hyphens and spaces
        text = re.sub(r"[^\w\s-]", " ", text)
        # Convert to lowercase
        text = text.lower()
        text = re.sub(r"\b\d+\b", "", text)
        text = re.sub(r"\bclass\b", "", text)
        text = re.sub(r"\bcare\b", "", text)
        text = re.sub(r"\bin\b", "", text)
        text = re.sub(r"\band\b", "", text)
        text = re.sub(r"\bthe\b", "", text)
        text = re.sub(r"\bfor\b", "", text)
        text = re.sub(r"\bwith\b", "", text)
        text = re.sub(r"\bfrom\b", "", text)
        text = re.sub(r"\bto\b", "", text)
        text = re.sub(r"\bunder\b", "", text)
        text = re.sub(r"\busing\b", "", text)
        text = re.sub(r"\bof\b", "", text)
        text = re.sub(r"\bno\b", "", text)
        text = re.sub(r"\binclude\b", "", text)
        text = re.sub(r"\bex\b", "", text)
        text = re.sub(r"\bexample\b", "", text)
        text = re.sub(r"\bclasses\b", "", text)
        text = re.sub(r"\bsearch\b", "", text)
        text = re.sub(r"\bscope\b", "", text)
        text = re.sub(r"\bproducts\b", "", text)

        # Standardize whitespace
        return " ".join(text.split())

    # Condition 3: Target market and goods/services overlap
    def target_market_and_goods_overlap(existing_gs, proposed_gs, threshold=0.65):
        embeddings1 = semantic_model.encode(existing_gs, convert_to_tensor=True)
        embeddings2 = semantic_model.encode(proposed_gs, convert_to_tensor=True)
        similarity_score = util.cos_sim(embeddings1, embeddings2).item()
        # st.write("Semantic Similarity Score:", similarity_score)
        if similarity_score >= threshold:
            return True

        # Normalize both strings
        existing_normalized = normalize_text(existing_gs)
        proposed_normalized = normalize_text(proposed_gs)

        # Split into words and lemmatize
        lemmatizer = WordNetLemmatizer()
        existing_words = {
            lemmatizer.lemmatize(word) for word in existing_normalized.split()
        }
        proposed_words = {
            lemmatizer.lemmatize(word) for word in proposed_normalized.split()
        }

        # Check for common words
        common_words = existing_words.intersection(proposed_words)
        # st.write("Common Words:", common_words)
        return bool(common_words)

    condition_3_satisfied = target_market_and_goods_overlap(
        existing_trademark["-_goods_&_services"], proposed_goods_services
    )

    # Clean and standardize the trademark -_status
    status = existing_trademark["-_status"].strip().lower()

    # Check for 'Cancelled' or 'Abandoned' -_status
    if any(keyword in status for keyword in ["cancelled", "abandoned"]):
        conflict_grade = "Low"
        reasoning = "The existing trademark -_status is 'Cancelled' or 'Abandoned.'"
    else:
        points = sum(
            [
                condition_1_satisfied,  # 1 point if any Condition 1 is satisfied
                condition_2_satisfied,  # 1 point if Condition 2 is satisfied
                condition_3_satisfied,  # 1 point if Condition 3 is satisfied
            ]
        )

        # Determine conflict grade based on points
        if points == 3:
            conflict_grade = "High"
        elif points == 2:
            conflict_grade = "Moderate"
        elif points == 1:
            conflict_grade = "Low"
        else:
            conflict_grade = compare_trademarks2(
                existing_trademark,
                proposed_name,
                proposed_class,
                proposed_goods_services,
            )

        if condition_1_satisfied:
            condition_1_details = []
            if condition_1A_satisfied:
                condition_1_details.append("Exact character-for-character match")
            if condition_1B_satisfied:
                condition_1_details.append("Semantically equivalent")
            if condition_1C_satisfied:
                condition_1_details.append("Phonetically equivalent")
            if condition_1D_satisfied:
                condition_1_details.append(
                    "First two or more words are phonetically equivalent"
                )
            if condition_1E_satisfied:
                condition_1_details.append(
                    "Proposed name is the first word of the existing trademark"
                )

        # Generate detailed reasoning for Condition 1
        if condition_1_satisfied:
            condition_1_reasoning = (
                f"Condition 1: Satisfied - {', '.join(condition_1_details)}."
            )
        else:
            condition_1_reasoning = "Condition 1: Not Satisfied."

        # Reasoning
        reasoning = (
            f"{condition_1_reasoning} \n"
            f"Condition 2: {'Satisfied' if condition_2_satisfied else 'Not Satisfied'} - Overlap in class numbers.\n"
            f"Condition 3: {'Satisfied' if condition_3_satisfied else 'Not Satisfied'} - Overlap in goods/services and target market."
        )

    # Return results
    return {
        "Trademark name": existing_trademark["-_trademark_name"],
        "Trademark -_status": existing_trademark["-_status"],
        "Trademark -_owner": existing_trademark["-_owner"],
        "Trademark class Number": existing_trademark["-_international_class_number"],
        "Trademark serial number": existing_trademark["-_serial_number"],
        "Trademark registration number": existing_trademark["-_registration_number"],
        "Trademark design phrase": existing_trademark["-_design_phrase"],
        "conflict_grade": conflict_grade,
        "reasoning": reasoning,
    }


def retrive_from_redis(redis_key):
    try:
        stored_data = redis_client.get(redis_key)
        if stored_data:
            details = json.loads(stored_data)
            # print(f"Details retrieved from Redis for {redis_key}: {details}")
            # st.write(f"Details retrieved from Redis for {redis_key}: {details}")
            return details
        else:
            print(f"No data found in Redis for key {redis_key}")
            return ""
    except Exception as e:
        print(f"Error retrieving data from Redis: {e}")
        return ""


if uploaded_files:
    new_files = []
    for uploaded_file in uploaded_files:
        if (
            uploaded_file.name
            not in [
                st.session_state.documents[doc_id]["name"]
                for doc_id in st.session_state.documents
            ]
            and uploaded_file.name not in st.session_state.removed_documents
        ):
            new_files.append(uploaded_file)

    for new_file in new_files:
        st.success(f"File Selected: {new_file.name}")
        pdf_bytes = new_file.read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        # st.write("hi2")
        page = doc[0]
        rect = page.rect
        height = 50
        clip = fitz.Rect(0, height, rect.width, rect.height - height)
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        # st.write("hi3")
        extracted_pages = []  # Array to store extracted text from each relevant page
        page_numbers = []  # Array to store corresponding page numbers
        extracted_pages2 = []  # Array to store text from all pages (optional)
        flag_uspto = False  # Flag to indicate USPTO Summary Page interval
        flag_state = False  # Flag to indicate State Summary Page interval

        for page_num, page in enumerate(doc, start=1):
            # Extract text with optional clipping
            text = page.get_text(clip=clip)
            extracted_pages2.append(text)

            # Check for interval boundaries
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text()
            if "USPTO Summary Page" in text:
                flag_uspto = True
            elif "ANALYST REVIEW −USPTO REPORT" in text:
                flag_uspto = False

            if "State Summary Page" in text:
                flag_state = True
            elif "ANALYST REVIEW −STATE REPORT" in text:
                flag_state = False

            # Store relevant text and page numbers for both intervals
            if flag_uspto or flag_state:
                extracted_pages.append(text)
                page_numbers.append(page_num)

        # st.write(extracted_pages2)
        # Initialize a list to store responses
        combined_responses = ""
        comparison_results = {
            "High": [],
            "Moderate": [],
            "Name-Match": [],
            "Low": [],
            "No-conflict": [],
        }
        st.write(extracted_pages)
        st.write(extracted_pages2)
        for extracted_text in extracted_pages:
            # st.write(extracted_text)
            prompt = f"""
                The task is to extract the name and associated page ranges in a structured JSON array format with each entry containing:
                - "name": The name of the entity (string).
                - "page-start": The first page number where the entity appears (integer)

                The data will be as below:
                '''
                1. ARRID EXTRA DRY Registered 3 CHURCH & DWIGHT CO., INC. 73−716,876 15

                2. ARRID EXTRA EXTRA DRY Registered 3 CHURCH & DWIGHT CO., INC. 78−446,679 18

                3. EXTRA RICH FOR DRY, THIRSTY HAIR Cancelled 3 NAMASTE LABORATORIES, L.L.C. 77−847,568 21
                
                4. GOOD LEAF Published 32, 33 DIAGEO NORTH AMERICA, INC. 90−829,139 89
                
                5. SHEAR GENIUS Registered 35, 44 SHEAR GENIUS OF FORT MO HAVE LLC 537444 AZ 225
                
                6. SHEAR GENIUS Registered 44 FABIO PAWLOS 1454759 NJ 226
                
                7. SHEAR GENIUS Registered 44 SHEAR GENIUS LLC 44423600 ND 227
                
                .
                .
                .
                '''
                
                It means that "ARRID EXTRA DRY" starts at page 15, "ARRID EXTRA EXTRA DRY" at page 16, "EXTRA RICH FOR DRY, THIRSTY HAIR" at page 21, "GOOD LEAF" at page 89, "SHEAR GENIUS" starts at page 225, "SHEAR GENIUS" at page 226, "SHEAR GENIUS" at page 227 and so on like that.
                Intelligently extract all entries of trademark name and their start page completely without leaving any entry from the given extracted text.
                The following text is extracted from a document:
                ---
                
                {extracted_text} 
                
                ---
                Just return the json without anyother additional text.
            """

            data = {
                "model": llm_model,
                "messages": [
                    {
                        "role": "system",
                        "content": "You are a helpful assistant that extracts details.",
                    },
                    {
                        "role": "user",
                        "content": prompt,
                    },
                ],
                "temperature": 0.0,
            }

            url = f"{azure_llm_endpoint}/openai/deployments/{llm_model}/chat/completions?api-version={llm_api_version}"
            llm_response = requests.post(
                url, headers=llm_headers, json=data, timeout=120
            )

            # Extract the response content
            response = (
                llm_response.json()
                .get("choices", [{}])[0]
                .get("message", {})
                .get("content", "")
                .strip()
            )

            record = response
            st_response = str(response)[7:-3]
            record = json.loads(st_response)
            # st.write(record)

            async def extract_trademark_details(document_chunk: str):
                try:
                    client = AzureOpenAI(
                        azure_endpoint=azure_llm_endpoint,
                        api_key=llm_api_key,
                        api_version="2024-10-01-preview",
                    )

                    messages = [
                        {
                            "role": "system",
                            "content": "You are a data extraction specialist proficient in parsing trademark documents.",
                        },
                        {
                            "role": "user",
                            "content": f"""
                            Extract the following details from the provided trademark document and present them in the exact format specified:  

                            - Trademark Name  
                            - Status  
                            - Serial Number  
                            - International Class Number (as a list of integers)
                            - Goods & Services (Goods and services are given after every international class, extract them intelligently as they may span over more than one page.)
                            - Owner  
                            - Filed Date (format: MMM DD, YYYY, e.g., Jun 14, 2024)  
                            - Registration Number  
                            - Design phrase

                            Instructions:  
                            - Return the results in the following format, replacing the example data with the extracted information:
                            - Ensure the output matches this format precisely.  
                            - Do not include any additional text or explanations.  

                            Document chunk to extract from:  
                            {document_chunk}  
                        """,
                        },
                    ]

                    loop = asyncio.get_event_loop()
                    response = await loop.run_in_executor(
                        None,
                        lambda: client.chat.completions.create(
                            model="gpt-4o-mini", messages=messages, temperature=0
                        ),
                    )

                    extracted_text = response.choices[0].message.content
                    details = {}
                    for line in extracted_text.split("\n"):
                        if ":" in line:
                            key, value = line.split(":", 1)
                            details[key.strip().lower().replace(" ", "_")] = (
                                value.strip()
                            )

                    return details

                except Exception as e:
                    logging.error(f"Error extracting trademark details: {e}")
                    return {"error": f"Error extracting trademark details: {str(e)}"}

            async def parallel_extraction():
                tasks = []
                for i in range(len(record)):
                    start_page = int(record[i]["page-start"]) - 1
                    if i == len(record) - 1:
                        end_page = start_page + 4
                    else:
                        end_page = int(record[i + 1]["page-start"])

                    document_chunk = "\n".join(extracted_pages2[start_page:end_page])
                    tasks.append(extract_trademark_details(document_chunk))

                return await asyncio.gather(*tasks)

            async def process_trademarks():
                extracted_details = await parallel_extraction()

                proposed_name = "GOOD HUMOR"
                proposed_class = "33"
                proposed_goods_services = "ALCOHOLIC BEVERAGES"

                for details in extracted_details:
                    if not details or "error" in details:
                        continue

                    comparision_result = compare_trademarks(
                        details, proposed_name, proposed_class, proposed_goods_services
                    )

                    conflict_grade = comparision_result.get("conflict_grade")
                    comparison_results[conflict_grade].append(comparision_result)

                # Create Word document

            asyncio.run(process_trademarks())

        # Create the document in memory
        doc = Document()
        for conflict_grade, results in comparison_results.items():
            count = len(results)
            doc.add_paragraph(f"{conflict_grade}: {count} entries")

            if results:
                doc.add_heading(conflict_grade, level=2)
                table = doc.add_table(rows=1, cols=5)
                header = table.rows[0].cells
                header[0].text = "Trademark Name and Class Number"
                header[1].text = "Trademark Status"
                header[2].text = "Serial/Registration Number"
                header[3].text = "Owner Name"
                header[4].text = "Design/Work"

                for result in results:
                    row = table.add_row().cells
                    row[0].text = (
                        f"{result['Trademark name']} (Class {result['Trademark class Number']})"
                    )
                    row[1].text = result["Trademark -_status"]
                    row[2].text = (
                        f"{result['Trademark serial number']} / {result['Trademark registration number']}"
                    )
                    row[3].text = result["Trademark -_owner"]
                    row[4].text = (
                        "Design" if result["Trademark design phrase"] else "Work"
                    )

        # Save the document to a BytesIO object
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        # Streamlit app logic
        st.title("Trademark Conflict Analysis")
        st.write("Download the trademark conflict analysis document below.")

        # Add a download button
        st.download_button(
            label="Download Analysis Document",
            data=output,
            file_name="Trademark_Conflict_Analysis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
